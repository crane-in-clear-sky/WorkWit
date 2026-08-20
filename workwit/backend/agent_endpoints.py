import io, json, re, traceback, asyncio, logging, time, os, subprocess, sys, zipfile, shutil, tempfile, uuid
from typing import List
import urllib.request, urllib.parse, html
import html as html_lib
from fastapi import FastAPI, UploadFile, File, Form, Request, Response, HTTPException
from fastapi.responses import FileResponse, JSONResponse, StreamingResponse
import db
import docx
from pypdf import PdfReader
from db import (
    init_db, get_user_by_token, create_session, delete_session, add_log,
    get_active, list_models, save_model, delete_model, activate, toggle_model_enabled,
    list_orgs, create_org, update_org, delete_org,
    list_departments, create_department, update_department, delete_department,
    list_users, create_user, update_user, delete_user, admin_count,
    get_user_permissions, set_user_permissions, has_permission, list_logs, list_logs_for_user,
    delete_logs_range,
    get_conn,
    list_tools, save_tool, inc_tool_calls, toggle_tool,
    save_skill, get_skill, get_skill_by_name, list_skills,
    delete_skill, toggle_skill, review_skill, inc_skill_calls, set_skill_visibility,
    update_skill, list_skill_versions, rollback_skill, clone_skill,
    install_skill, uninstall_skill,
    get_active_caps, add_active_caps, clear_active_caps,
    _SKILL_NAME_RE,
    MASK,
    PERMS, PERM_LABELS,
)
from agent import run_agent, run_agent_with_retry, resolve_session_tools
import sandbox

from core import (ToolContext, _model_params, _sse, _detect_generated_files,
                _ext_of, AGENT_TASKS, extract_text)
from builtin_tools._shared import _user_dir, _file_belongs_to
from auth import require_perm, client_ip
from db import (get_active, get_conn, list_tools, list_skills, add_log,
               inc_tool_calls, inc_skill_calls,
               get_user_memory, delete_user_memory, list_tasks,
               create_automation, list_automations, get_automation,
               update_automation, delete_automation, list_automation_runs,
               ensure_session, list_sessions, update_session_title,
               delete_session, get_session_messages,
               get_user_profile, save_user_profile)
from agent import run_agent, run_agent_with_retry, resolve_session_tools
from llm_adapter import create_client, ModelCaps
from search import search_web, fetch_page_text
from tools_build import (build_meta_tools, build_session_tools,
                         build_session_skill_tools, _build_skill_constraints,
                         _build_method_prompt, _parse_method_markers,
                         _narrow_tools_by_whitelist, _has_lexical_hit,
                         _apply_skip_skill)

try:
    import mcp_client
except Exception:
    mcp_client = None

import automation_runner

from fastapi import APIRouter
router = APIRouter()


async def _produce_plan(client, model_name, params, question, catalog_text, method_text):
    """P1⑥ 规划阶段：生成分步执行计划文本（不执行任何工具）。

    用一次非流式 LLM 调用产出计划，使复杂/多步任务在动手前先经用户确认，前置纠偏。
    返回计划字符串；异常时返回带错误说明的字符串（Gen() 仍会把该文本下发给前端）。
    """
    sys_prompt = (
        "你是一个严谨的任务规划助手。用户会给出一项任务，请基于下方「可用能力」"
        "输出一份**分步执行计划**。要求：\n"
        "1. 用编号列表，每步写清：要做什么、用到哪个工具/技能（从可用能力中选取，写其名称）、预期产出什么。\n"
        "2. 若任务缺少关键业务信息（分析对象、行业、范围、竞品、具体需求等），"
        "在计划末尾单独用「⚠️ 需向用户澄清：」列出待提问项，不要凭空假设。\n"
        "3. 只输出计划本身，**不要执行任何工具**，也不要写代码。\n"
        "4. 计划应具体、可执行，避免空泛；步骤数控制在 3–8 步。"
    )
    if catalog_text:
        sys_prompt += "\n\n【可用能力（工具/技能目录）】\n" + catalog_text
    if method_text:
        sys_prompt += "\n\n" + method_text
    try:
        resp = await asyncio.wait_for(
            asyncio.to_thread(
                client.chat.completions.create,
                model=model_name,
                messages=[{"role": "system", "content": sys_prompt},
                          {"role": "user", "content": question}],
                **params,
            ),
            timeout=60,
        )
        return (resp.choices[0].message.content or "").strip() or "（模型未返回计划内容）"
    except Exception as e:
        return f"（生成计划时出错：{type(e).__name__}: {e}）"


_GREETING_TOKENS = {
    "你好", "您好", "在吗", "在么", "在不在", "有人吗", "hi", "hello", "hey",
    "哈喽", "嗨", "早上好", "中午好", "下午好", "晚上好", "晚安",
    "谢谢", "感谢", "多谢", "辛苦了", "再见", "拜拜", "回头聊",
}


def _is_greeting_or_chitchat(text):
    """判断是否为纯闲聊 / 仅含问候语。这类消息不应触发技能安装推荐。

    规则（去标点与空白后）：整体去掉所有问候语词后，剩余仅语气词或为空 → 闲聊；
    其余（含明确任务动词、较长语句）视为有意图，不拦截。
    """
    if not text:
        return True
    s = re.sub(r"[\s，。！？、；：,.!?;:\"'\-—…（）()\[\]【】~@#$%^&*=+\\|/]",
               "", (text or "").lower())
    if not s:
        return True
    residual = s
    for g in _GREETING_TOKENS:
        residual = residual.replace(g, "")
    if residual in ("", "啊", "呀", "呢", "吧", "哦", "啦", "嘛", "哈", "额", "诶", "呵", "嘞"):
        return True
    return False


_UPLOAD_ROOT = os.environ.get("UPLOAD_ROOT", "/app/data/uploads")
_UPLOAD_EXTS = (".txt", ".docx", ".pdf", ".md", ".csv", ".json", ".py", ".log")
_UPLOAD_MAX_MB = 20


@router.post("/api/agent/upload")
async def agent_upload(request: Request, file: UploadFile = File(...)):
    """上传落盘端点：把用户附件（尤其 docx/pdf 等前端 FileReader 抽不出文本的二进制文档）
    保存到服务器数据目录，返回绝对路径。前端拿到 path 后随 chat 请求以 attachments[].path 传入，
    后端再读文件抽文本注入上下文——解决「输入资料没进上下文」导致的凭空编造。"""
    u = require_perm("agent", request)
    name = (file.filename or "attachment").strip()
    ext = os.path.splitext(name)[1].lower()
    if ext not in _UPLOAD_EXTS:
        return JSONResponse(status_code=400, content={
            "error": "不支持的文件类型：%s（仅支持 txt/docx/pdf/md/csv/json/py/log 等文本类）" % (ext or "无扩展名")})
    data = await file.read()
    if not data:
        return JSONResponse(status_code=400, content={"error": "文件为空"})
    if len(data) > _UPLOAD_MAX_MB * 1024 * 1024:
        return JSONResponse(status_code=400, content={"error": "文件过大（上限 %dMB）" % _UPLOAD_MAX_MB})
    # 用户级隔离：上传附件按 <user_id>/ 子目录落盘，避免平铺导致跨用户越权
    _updir = _user_dir(u["id"], None, "uploads")
    try:
        os.makedirs(_updir, exist_ok=True)
    except Exception as e:
        return JSONResponse(status_code=500, content={"error": "无法创建上传目录：%s" % e})
    safe = re.sub(r"\W+", "_", os.path.splitext(name)[0])[:40] or "attachment"
    fname = "%s_%s%s" % (safe, uuid.uuid4().hex[:8], ext)
    dst = os.path.join(_updir, fname)
    try:
        with open(dst, "wb") as f:
            f.write(data)
    except Exception as e:
        return JSONResponse(status_code=500, content={"error": "文件落盘失败：%s" % e})
    return {"ok": True, "path": dst, "name": name, "ext": ext, "size": len(data)}


@router.post("/api/agent/chat")
async def agent_chat(request: Request, payload: dict):
    """自主智能体对话：接收用户目标，跑 ReAct 循环自动调用工具，SSE 流式返回规划与结果。"""
    u = require_perm("agent", request)
    active = get_active("chat") or {}
    # 支持前端下拉指定主推理模型（仅限 role=chat 且已启用，避免越权选到其他角色/禁用模型）
    sel_id = payload.get("model_id")
    if sel_id:
        try:
            cid = int(sel_id)
            if cid > 0:
                c = get_conn()
                row = c.execute(
                    "SELECT * FROM models WHERE id=? AND role='chat' AND enabled=1", (cid,)).fetchone()
                c.close()
                if row:
                    active = dict(row)
        except Exception:
            pass
    base = (active.get("base_url") or "").strip()
    key = (active.get("api_key") or "").strip()
    name = (active.get("model_name") or "").strip()
    if not base:
        return JSONResponse(status_code=200, content={
            "error": "未配置主推理模型",
            "message": "请到「系统管理 → 模型配置」启用一个 chat 模型"})
    # P4·多模型兼容：用统一适配层创建客户端（任意 OpenAI 兼容端点），超时由模型能力声明驱动
    caps = ModelCaps(active)
    client = create_client(base, key or "not-needed", caps.client_timeout())
    params = _model_params(active)
    params.setdefault("temperature", 0.3)

    question = (payload.get("message") or "").strip()
    if not question:
        return JSONResponse(status_code=400, content={"error": "请输入内容"})
    history = payload.get("history") or []
    # 工具调用模式：仅控制「工具匹配」子步骤的粒度，不影响顶层路由（简单任务→LLM自身能力 / 复杂任务→技能优先）
    # full（A·全量注入：命中能力全量注入由 LLM 自选，默认）/ keyword（B·智能精选：路由 LLM 语义精选工具子集）
    tool_mode = (payload.get("tool_mode") or "full").strip().lower()
    if tool_mode not in ("keyword", "full"):
        tool_mode = "full"
    # P1⑥ Plan 模式：chat=默认直接执行；plan=先出分步计划待确认；execute=按已确认计划执行
    mode = (payload.get("mode") or "chat").strip().lower()
    if mode not in ("chat", "plan", "execute"):
        mode = "chat"

    # 联网模式：检索资料拼接进用户问题（沙箱外主进程联网，容器网络默认放通）
    web_hits = []
    web_engine = ""
    if bool(payload.get("web")):
        try:
            # 同步 IO 扔到线程池，并用 wait_for 卡死总时间，避免阻塞事件循环
            web_hits = await asyncio.wait_for(
                asyncio.to_thread(search_web, question, 5),
                timeout=12
            )
        except asyncio.TimeoutError:
            web_hits = []
        except Exception:
            web_hits = []
        if web_hits:
            web_engine = web_hits[0].get("engine", "")
            # 深读：摘要常缺具体数值（温度/价格/日期），抓取 Top N 正文补足
            try:
                n_deep = int(os.environ.get("WEB_FETCH_PAGES", "2"))
            except Exception:
                n_deep = 2
            pages = {}
            if n_deep > 0:
                try:
                    targets = [h["url"] for h in web_hits[:n_deep] if h.get("url")]
                    texts = await asyncio.wait_for(
                        asyncio.gather(*[asyncio.to_thread(fetch_page_text, u) for u in targets],
                                       return_exceptions=True),
                        timeout=10
                    )
                    for src, t in zip(targets, texts):
                        if isinstance(t, str) and len(t) > 80:
                            pages[src] = t
                except Exception:
                    pass
            lines = []
            for i, h in enumerate(web_hits):
                lines.append(f"[{i+1}] {h['title']}\n    来源：{h.get('url', '')}"
                             f"\n    摘要：{h.get('snippet', '')}")
                body = pages.get(h.get("url", ""))
                if body:
                    lines.append(f"    正文节选：{body}")
            question = (question +
                        f"\n\n【联网检索资料（引擎：{web_engine}，当前时间 {time.strftime('%Y-%m-%d %H:%M')}）】\n"
                        + "\n".join(lines) +
                        "\n\n请优先依据上述实时资料作答，并在结论中标注信息来源；"
                        "若资料互相矛盾或明显过期，请指出并结合你的知识判断。")
    # 附件：①文本类文件由前端 FileReader 内联 content；②docx/pdf 等二进制文档前端先上传落盘、
    #   随请求带 path，此处后端读文件抽文本注入上下文（解决「输入资料没进上下文」的凭空编造）。
    atts = payload.get("attachments") or []
    if isinstance(atts, list) and atts:
        parts = []
        for a in atts:
            if not isinstance(a, dict):
                continue
            _name = a.get("name", "附件")
            _content = (a.get("content") or "").strip()
            _path = (a.get("path") or "").strip()
            if _content:
                parts.append("--- 文件：%s ---\n%s" % (_name, _content))
            elif _path:
                try:
                    if not os.path.isfile(_path):
                        parts.append("--- 文件：%s ---\n（文件不存在：%s）" % (_name, _path))
                        continue
                    with open(_path, "rb") as _f:
                        _raw = _f.read()
                    _txt = extract_text(_raw, os.path.basename(_path))
                    if _txt and _txt.strip():
                        parts.append("--- 文件：%s ---\n%s" % (_name, _txt.strip()))
                    else:
                        parts.append("--- 文件：%s ---\n（文档内容为空或无法解析）" % _name)
                except Exception as _e:
                    parts.append("--- 文件：%s ---\n（读取失败：%s）" % (_name, _e))
        if parts:
            question = question + "\n\n【附件内容】\n" + "\n".join(parts)

    messages = [{"role": "system", "content":
        "你是一个自主规划的智能助理。"
        "【🔧 工具使用与创建规范】"
        "**优先使用已有工具**：执行任务前先查看已注入的工具列表，如果有能完成任务的工具，直接调用它。"
        "**禁止绕过确认直接创建**：你不得直接调用 create_tool 创建新工具。如果需要新工具，必须走「提议→用户确认→创建」流程（见下方缺工具兜底）。"
        "**禁止重复创建**：全局工具库中可能已有大量历史工具（Word生成、文档输出、PPT等）。"
        "在提议/创建新工具前，必须先调用 list_available_tools（可用关键词过滤）检查是否已存在功能相似的工具。"
        "如果已有相似工具，直接使用它，不要创建新的。服务端有去重检测，重复创建会被拒绝。"
        "**创建失败后立即停止**：如果 create_tool 返回错误（如「功能相似的工具已存在」），立即停止尝试，改用系统返回的已有工具或自身能力回答。严禁换名字再调。"
        "**创建成功后禁止再创建（绝对铁律）**：create_tool 返回「✅ 已创建成功」后，你必须在下一步立即调用该新工具执行任务。"
        "**严禁因「对结果不满意」而再次 create_tool**——如果已有工具已执行并产出文件（如 PPT/Word/Excel），即使内容不完美，也视为任务完成。"
        "如需改进质量，应调整传给该工具的参数（如更详细的 content/text），而非重新创建工具。每次对话中 create_tool 最多调用 1 次，超出会被系统强制拦截。"
        "【完整输出原则】在最终回复中，你必须**完整输出任务结果的所有内容**——不要只给摘要、概要或「报告已生成，涵盖XXX」之类的简略描述。"
        "如果生成了文档（Word/PPT/Excel 等），你必须在文本回复中**完整呈现该文档的全部内容**（包括所有章节、数据、分析细节），让用户在聊天区就能看到完整结果，而不仅仅是一个文件下载链接。"
        "【方法论优先原则】如果本次对话已注入了【方法论技能 / 专家流程】（下方会有该段落），你必须**严格按其框架/流程执行任务**。执行时遵循「缺参必问」：**若框架任一步骤所需的关键业务参数（分析对象/产品、行业、范围、竞品、具体需求等）用户尚未在对话中提供，你必须先通过提问向用户澄清，待其回复后再继续该步骤——严禁凭空假设、自行决定或编造这些关键输入。** 例如竞品分析类任务，若用户未说明分析哪款产品、哪个行业、关注哪些竞品，必须先逐一询问，拿到答案后再按框架分析，不得直接编造报告。其余步骤按框架收集信息、分析、再输出结论。只有框架明确要求或用户主动询问时，才调用 list_available_tools / list_available_skills 等元工具。不要在已注入方法论的情况下第一时间去探索工具列表。"
        "**优先复用已有技能**：若系统已注入或已安装相关方法论技能，直接按其框架执行，不要新建技能。"
        "如需创建新技能（例如任务中途确实需要一种全新的思考框架），你可以调用 create_skill，但**创建前系统会自动查重**——"
        "若检测到与已有技能（含技能广场公开技能、你本人私有技能）名称相同或触发词高度重叠，会拒绝重复创建并提示你直接使用已有技能，请勿换名字反复重试。"
        "若执行中发现某个具体**工具能力**（如更好的 Word 排版工具）缺失，应走 [PROPOSE_TOOL] 提议创建**新工具**而非新技能。"
        "对于没有注入方法论技能的任务，当任务可以用已有工具完成时（如总结文本、分析合同风险、评估简历、生成 PPT、计算），"
        "请主动调用合适的工具；多步任务请分步调用，最后用中文**完整、详细地呈现全部结果内容**（包括分析过程、数据细节、完整结论），不要省略或概括。"
        "【技能生态】系统中有一个「技能广场」，存放着各类方法论技能和代码工具技能。用户可以安装技能，安装后你就能在任务中自动使用它们。"
        "你可以随时调用 list_available_skills 工具查看当前已安装的可用技能列表（支持关键词过滤），了解系统已有的能力。"
        "在创建新技能之前，建议先调用 list_available_skills 检查是否已有类似技能，避免重复创建。"
        "【全局工具】除技能外，系统还在「全局工具库」中注册了一系列内置能力（如 summarize_text、calculator、analyze_contract_risk、screen_candidate、generate_ppt、generate_word 等）。"
        "它们由系统按本次任务自动检索注入，遵循「技能优先」原则：方法论技能优先作为思考框架注入系统提示；复杂任务在技能框架内，再按当前工具模式注入相关工具——"
        "全量注入=命中的相关能力（工具+代码技能）全量注入、由你用 tool_call 自选；智能精选=路由 LLM 语义精选子集。简单对话/简单任务则不注入任何工具，直接由你自身能力作答。无需用户安装。**执行任务时优先使用已注入的现有工具，并严格遵循已注入的方法论技能框架。**"
        "你可以调用 list_available_tools 工具查看全局工具库中已注册的全部工具列表（支持关键词过滤）。"
        "【向用户确认·AskUserQuestion】执行任务过程中若遇到「必须和用户确认才能继续」的关键参数/方案选择（如缺失的业务对象、互斥选项），"
        "不要凭空假设，应调用 ask_user 工具提出问题与候选选项；系统会弹出对话框等你回复，回复后任务自动继续。"
        "【任务拆解·Task*】对于长程、多步骤的复杂任务，应先用 task_create 拆出子任务、用 task_update 维护状态（pending→in_progress→completed）、"
        "用 task_list 查看整体进度——这能让长任务的状态与进展被持续追踪、不因轮次/步骤多而丢失。"
        "**如果确实找不到能完成任务的工具，不要直接创建，也不要让用户手动操作——通过 [PROPOSE_TOOL] 标记提议创建，系统会请用户确认后自动生成。**"
        "当用户询问「系统有什么工具/能力」时，可同时调用 list_available_skills 与 list_available_tools 给出完整回答。"
        "【📄 内置生成工具调用铁律（generate_ppt / generate_word）——已支持表格/图表/排版，无需创建替代工具】"
        "**调用 generate_ppt 时：topic 填主题（作为封面标题），content 参数必须填入你要写入 PPT 的全部正文内容（用 --- 分页、或 # 标题分页、或空行分段）。严禁只传 topic 不传 content——那样生成的 PPT 只有封面和占位文字没有实质内容。**"
        "**调用 generate_word 时：content 参数必须填入你要写入文档的全部正文内容（支持 Markdown 表格语法 | 列 | 列 | ，会自动渲染为带边框、对齐、表头样式的真实 Word 表格），title 填文档标题。严禁传空内容。generate_word 已完整支持表格/标题/列表/加粗等排版，不需要为此创建新工具。**"
        "**正确示例**：generate_ppt(topic='季度汇报', content='---\n# 销售业绩\n- Q1增长23%\n- Q2增长31%\n\n# 产品进展\n- 完成V2.0开发\n- 上线新功能')**"
        "若用户要求「创建一个技能（用来指导某种做法/思考方式）」，请调用 create_skill 工具："
        "创建「方法论技能」（skill_type=method，在 instructions 参数中直接写出中文提示词/流程，不写 Python 代码；"
        "若确实未写，系统会自动生成兜底版本，但请尽量亲笔撰写以保证技能质量）。"
        "若用户要求「创建一个工具（写一段可运行的 Python 代码实现某功能）」，请改用 create_tool 工具："
        "在 code 参数中写出 Python 代码（必须定义 def run(a): ... 且仅做纯数据处理），保存为私有工具存于「全局工具库」，沙箱执行。"
        "【⚠️ 文档生成工具的代码规范（Word/PPT/Excel 必须严格遵守）】"
        "如果创建的是文档输出类工具（如 Word/PPT/Excel/PDF 生成），你的 code 代码**必须**满足以下要求，否则生成的文档将为空或无法使用："
        "(1) run(a) 的参数 a 是一个 dict，调用者传入的实际内容可能在 a['text']、a['content']、a['input'] 等键下。"
        "**你必须用 a.get('text', '') or a.get('content', '') or a.get('input', '') 的方式兼容取值，不要只写 a['text'] 这样单一取值。**"
        "(2) 取到的内容就是需要写入文档的全部正文，**必须逐段/逐行写入文档**（用 add_paragraph/add_slide 等），不要遗漏或省略。"
        "(3) 文档保存路径必须是绝对路径（如 /tmp/output.docx），且 return 字符串中必须包含该完整路径（格式：「已生成：/tmp/output.docx」），否则前端无法展示下载按钮。"
        "(4) 正确的 Word 生成工具代码模板如下（请严格参考此结构）："
        "```python"
        "from docx import Document"
        "from docx.shared import Pt, RGBColor"
        "import os"
        "def run(a):"
        "    # 兼容多种参数名取内容"
        "    text = str(a.get('text', '') or a.get('content', '') or a.get('input', '') or '')"
        "    if not text.strip():"
        "        return '错误：未提供文档内容'"
        "    doc = Document()"
        "    # 写入标题"
        "    title = str(a.get('title', '') or a.get('subject', '') or '文档')"
        "    doc.add_heading(title, level=0)"
        "    # 按段落分割并逐段写入（保留原始格式）"
        "    for para in text.split('\\n'):"
        "        if para.strip():"
        "            doc.add_paragraph(para.strip())"
        "    # 保存并返回路径"
        "    out_path = '/tmp/output.docx'"
        "    doc.save(out_path)"
        "    size = os.path.getsize(out_path)"
        "    return f'文档已生成（{size} 字节）\\n已生成：{out_path}'"
        "```"
        "**(5) 对于 PPT/Excel 同理：必须把 a 中的实际内容完整写入每一页/每个单元格，不要生成空壳文件。）**"
        "并务必生成 3-5 个中文触发词填入 trigger_words 参数（如合同审查类填「合同审查,审查要点,合同风险」），便于智能体通过关键词检索命中该技能/工具；"
        "同时务必用一句话填写 when_to_use 参数（描述「用户说什么 / 处于什么场景时该用此技能」，如「用户要求审查合同、给出风险清单时」），供路由 LLM 语义匹配，与触发词互补提升命中准确率；"
        "同时务必提供 display_name 参数（中文展示名，将显示为「中文（英文）」，如「合同审查指引」，不要用英文、不要留空）；"
        "创建的技能/工具均为私有（仅创建者本人可见、可用），之后你自己及该用户的相关任务会自动检索并注入执行。"
        "【缺技能兜底·人机协作】**重要前提：只有当本次对话完全没有注入任何【方法论技能 / 专家流程】时，才允许触发以下流程。"
        "如果 system prompt 中已有「【方法论技能 / 专家流程」段落（即已注入了 method 类型技能），说明系统已为你提供了思考框架，你必须直接按其执行，严禁再输出 [PROPOSE_METHOD] 标记或建议保存新技能——这会造成重复。**"
        "仅在没有注入任何方法论技能的前提下，当你判断当前任务涉及多步骤复杂流程、需要明确思考框架（如审查/调研/方案撰写/规划类任务）时："
        "(1) 第一次回复时，先用一段中文写出你建议的「思考步骤/做法框架」草稿（写清分几步、每步做什么、最终交付什么），"
        "然后在回复末尾追加标记 [PROPOSE_METHOD] 草稿全文 [/PROPOSE_METHOD]（标记内不要含其他内容），系统会弹出对话框让用户确认/编辑。"
        "(2) 用户确认后，会以「[用户已确认思考框架]：...」开头发起新一轮对话——此时请按确认后的框架执行任务，不要重新生成框架。"
        "(3) 任务执行完毕、给出最终结论后，**仅当你在第(1)步确实输出了 [PROPOSE_METHOD] 标记**（即当时没有已注入的方法论技能），"
        "才请在最终回复末尾追加标记 [SUGGEST_SAVE_SKILL] 建议中文展示名 | 建议英文标识符 [/SUGGEST_SAVE_SKILL]。"
        "如果本次已有方法论技能注入、你没有输出过 [PROPOSE_METHOD]，则绝对不要输出 [SUGGEST_SAVE_SKILL]——否则会创建重复技能。"
        "系统会弹出「保存为技能」对话框。用户同意后会以「是，请把本次思考框架保存为技能，display_name=XX, name=YY」发起新一轮对话，"
        "此时请调用 create_skill 工具（skill_type=method，instructions 填写用户确认过的思考框架全文，display_name/name 用用户给的值）保存为私有技能。"
        "【缺工具兜底·人机协作】当任务需要某个工具（如 Word 文档生成、PDF 转换、Excel 处理等）但全局工具库中没有时："
        "**第一步（必须，仅限 1 次）**：调用 list_available_tools 工具（可用关键词过滤，如「Word」「文档」「生成」），确认确实没有可用的相关工具。"
        "**重要：list_available_tools 最多只调 1 次**——调用后根据返回结果立即决策：有合适工具就直接用、没有就走第二步。绝对不要反复调用它（系统有频率限制，超限会被拦截）。"
        "**第二步（必须）**：确认无匹配工具后，你必须在最终回复末尾追加标记 [PROPOSE_TOOL] 工具用途一句话描述 [/PROPOSE_TOOL]（标记内只放描述，不含代码）。"
        "系统会弹出对话框让用户确认是否生成私有工具。**这是唯一的创建新工具的合法路径，严禁绕过用户确认直接调用 create_tool。**"
        "**补充（已有工具但能力不完整）：** 若你已调用某个已存在的工具，但其输出表明它无法完成当前子任务"
        "（返回错误、空结果、或明显缺少所需能力），不要反复调用它凑合——请在最终回复末尾追加 [PROPOSE_TOOL] 标记，"
        "提议创建一个功能更完整的工具（用户确认后创建并使用）。这条规则专门覆盖「库里同名/相似工具能力不够用」的情形。"
        "(1) 用户确认后，会以「是，请创建该工具并立即执行任务」发起新一轮对话——此时请按以下顺序执行："
        "  ① 先调用 create_tool 工具（在 code 参数中写出 Python 代码，定义 def run(a): ... 且仅做纯数据处理）；"
        "  **② create_tool 返回成功后（会提示「✅ 工具已创建成功」），你必须在下一步立即调用刚创建的工具来执行任务**——不要再去 list_available_tools、不要再次 create_tool、不要做其他无关操作。直接调新工具，把用户要处理的内容传给 text/content 参数。"
        "**关键：如果生成的是文档类工具（Word/PPT/Excel），代码必须从参数 a 中正确提取内容并完整写入文档（用 a.get('text','') or a.get('content','') 兼容取值），不要生成空壳文档。**"
        "(2) 若用户拒绝创建工具，你仍可在沙箱中执行一段**临时代码**来完成任务——调用 run_temp_code 元工具"
        "（传入 code 与待处理内容 text/content），其执行结果可用、产物可下载，但**不会保存为工具**（用户拒绝持久化）。"
        "若连临时代码也无法满足，则用你自身能力尽量完成，或告知用户该功能暂不可用。"
        "**绝对禁止的行为**：不要告诉用户「建议您手动注册」「请在全局工具库中手动创建」之类的话——系统有完整的自动创建流程，你必须通过 [PROPOSE_TOOL] 触发它。"
        "重要：对于天气、百科、常识、定义、闲聊等通用知识类问题，优先直接用你的知识回答，无需调用任何工具，也不要触发缺技能/缺工具兜底流程。"}]
    # 历史注入：窗口内逐字 + 窗口外压缩为摘要（P2·compaction），避免长对话早期任务背景丢失
    for m in _compact_history(history, keep=12):
        messages.append(m)
    messages.append({"role": "user", "content": question})

    # 全局工具库 + 用户可用技能 检索注入：解析本次任务的能力缺口，从库中挑选相关工具/技能
    # 临时注入会话，避免把所有能力常驻上下文；路由 LLM 无命中则保底全部注入。
    # 传 for_user_id：用户私有工具（is_user_created=1, scope=private）也进 library，对自己可见可用
    tools_lib = list_tools(for_user_id=u["id"])
    skills_lib = list_skills(for_user_id=u["id"], usable_only=True, with_code=True)
    tool_names = {t["name"] for t in tools_lib}
    skill_names = {s["name"] for s in skills_lib}
    # 用户在前端「技能选择器」中强制指定的技能：无论路由 LLM 是否命中，都强制注入为约束。
    # 仅接受用户可见可用的技能（已发布/已安装/本人私有），避免越权注入他人私有技能。
    _sel_skill_def = None
    _sel_skill_id = payload.get("selected_skill_id")
    if _sel_skill_id not in (None, "", 0):
        try:
            _sid = int(_sel_skill_id)
            _cand = get_skill(_sid, with_code=True)
            if _cand and any(s.get("id") == _sid for s in skills_lib):
                _sel_skill_def = _cand
        except (ValueError, TypeError, Exception):
            _sel_skill_def = None
    # 层 1 · method 技能增强注入：关键词命中 question 的方法论技能强制抽出注入 system prompt。
    # 这是「增加」信号（非闸门）——仅确保强相关 method 必注入，不剥夺 LLM 对其他能力的判断。
    # 真正的跨轮防脱落由下方 session_active_caps 持久化保证（一旦激活持续在场，不依赖关键词）。
    method_hit_defs = [s for s in skills_lib
                       if s.get("skill_type") == "method"
                       and _has_lexical_hit(question, [s])]
    # 路由 LLM 看的 library：剔除「已命中」的 method 技能（避免重复考虑），
    # 保留「未命中」的 method 技能——交给路由 LLM 基于语义判断要不要注入
    method_hit_names = {s["name"] for s in method_hit_defs}
    skills_lib_for_router = [s for s in skills_lib
                             if s.get("skill_type") != "method"
                             or s["name"] not in method_hit_names]
    library_for_router = tools_lib + skills_lib_for_router

    # 会话标识：前端稳定传递 session_id（同一连续对话复用，开新对话换新）；
    # 旧前端不传时降级使用 task_id（单轮，无跨轮持久，向后兼容）。
    task_id = payload.get("task_id") or os.urandom(12).hex()
    session_id = payload.get("session_id") or task_id
    ctx = ToolContext(client, name, params, user=u)
    ctx.session_id = session_id   # 供子 Agent 派生独立子会话（P1⑤）
    # P2⑪ 对话消息落库 + 会话记录建立：供跨对话历史检索与会话管理（按窗口删除/续聊）
    try:
        db.save_conversation_message(session_id, u["id"], "user", question)
        db.ensure_session(u["id"], session_id, question)
    except Exception as _e:
        print(f"[P2⑪] 对话落库跳过: {_e}")
    mx = payload.get("max_create_skills")
    if isinstance(mx, int) and mx > 0:
        ctx.max_create_skills = mx
    meta = build_meta_tools(ctx)
    # 元工具（如 create_skill）始终可用，单独并入（不参与检索）
    meta_defs = [{"name": t["name"], "description": t["description"], "kind": "meta"} for t in meta]
    # 注：create_skill 的剔除统一在 method_text 计算之后（下方），避免此处用 method_hit_defs（仅词面匹配）
    # 作为条件信号时漏掉「路由 LLM 语义匹配」选中的方法论技能。

    # ────────────────────────────────────────────────────────────────────────
    # P1 新匹配：常驻目录(全量注入 system prompt) + 路由 LLM 语义预展开(top-k)
    #           + 会话已激活能力持久(只增不减,防脱落) —— 三者共同替代旧「词面闸门」。
    # 匹配信号只负责「增加」激活能力，绝不负责「移除」：即便路由漏选，
    # 已激活能力因 session_active_caps 持久化而持续在场 → 根除多轮「已加载技能脱落」。
    # ────────────────────────────────────────────────────────────────────────
    if tool_mode == "full":
        # A 模式·全量注入：相关能力（工具 + code 技能）全量注入，由 LLM 用 tool_call 自选
        router_sel = [t for t in library_for_router if t.get("skill_type") != "method"]
    else:
        # B 模式（默认）·路由 LLM 语义预展开 top-k：基于「当前任务 + 常驻目录」挑选候选（非闸门）
        router_sel = await resolve_session_tools(question, library_for_router, client, name, params)

    # 会话已激活能力补回（只增不减）：把上一轮已加载、本轮路由未选中的能力续上
    # 2026-08-19 隔离加固：传 user_id，避免 session_id 撞车时继承对方的能力集
    active_names = get_active_caps(session_id, user_id=u["id"])
    selected = list(router_sel)
    sel_names = {d["name"] for d in selected}
    for t in library_for_router:
        if t["name"] in active_names and t["name"] not in sel_names:
            selected.append(t)
            sel_names.add(t["name"])

    # 元工具常驻 + method 关键词增强（由 question 判定，属「增加」信号，非闸门）
    selected = selected + meta_defs + method_hit_defs

    # 用户强制指定技能：强制并入 selected，保证其进入 skill_defs
    # → method 技能注入 system prompt 框架 / code 技能包装为工具 / 约束收窄工具面
    # （不依赖路由 LLM 是否选中，满足「用户选了就要按它约束执行」）
    if _sel_skill_def and not any(d.get("id") == _sel_skill_def.get("id") for d in selected):
        selected.append(_sel_skill_def)

    # 持久化本轮激活（只增不减）：本轮匹配到的全部能力记入会话激活集，供后续轮次续接
    # 2026-08-19 隔离加固：传 user_id，激活集按用户隔离
    add_active_caps(session_id, [d["name"] for d in selected], user_id=u["id"])

    # 场景二(A)：命中带 skip_skill 的全局工具 → 真正跳过 Skill 匹配，仅保留工具 + 元工具
    selected = _apply_skip_skill(selected, tool_names)

    tool_defs = [d for d in selected if d["name"] in tool_names]
    skill_defs = [d for d in selected if d["name"] in skill_names]

    # 场景四：被命中的技能若声明了业务规则/工具清单，注入上下文并收窄工具选择面
    constraint_text, allowed_union = _build_skill_constraints(skill_defs)
    # 方法论技能（skill_type=method）：提示词/流程注入 system prompt，约束思考方式，不包装成工具
    method_text = _build_method_prompt(skill_defs)
    extra_sys = ""
    if constraint_text:
        extra_sys += constraint_text
    if method_text:
        extra_sys += ("\n\n" if extra_sys else "") + method_text
    # 用户在前端「技能选择器」强制选定的技能：明确告知 LLM 这是硬性约束，
    # 必须在该技能框架/约束下执行，不得脱离或另起炉灶、也不要再提议创建类似技能。
    if _sel_skill_def:
        _dn = _sel_skill_def.get("display_name") or _sel_skill_def.get("name") or "指定技能"
        extra_sys += ("\n\n" if extra_sys else "") + (
            "【用户强制指定技能】本次任务用户明确选择了技能「%s」，你必须严格按该技能定义的"
            "框架 / 流程 / 约束执行任务，不得脱离该技能自行另起炉灶，也不要再提议创建功能类似的技能。"
            % _dn)
    # 常驻能力目录（全量轻量，每轮注入）：让 LLM 始终看到「完整可用能力菜单」原生推理，
    # 而非只看到路由预展开的 top-k 子集——这是「匹配=LLM 推理式」的载体（§1.1）。
    catalog_text = _render_capability_catalog(tools_lib, skills_lib)
    if catalog_text:
        extra_sys += ("\n\n" if extra_sys else "") + catalog_text
    # 用户跨会话长期记忆（对标 WorkBuddy「云记忆自动注入」）：每轮读该用户记忆常驻注入，
    # 让智能体「记得」用户偏好 / 项目背景 / 双方约定，跨新对话仍生效。异常降级为空（绝不终止任务）。
    try:
        _mems = get_user_memory(u["id"]) if u else []
    except Exception:
        _mems = []
    if _mems:
        _lines = []
        for _m in _mems:
            _type = _m.get("mem_type") or "note"
            _lines.append(f"- [{_type}] {_m.get('content', '')}")
        extra_sys += ("\n\n" if extra_sys else "") + (
            "<user_long_term_memory>\n以下是该用户的跨会话长期记忆（下次新对话也会自动带上，请据此调整表述与行为）：\n"
            + "\n".join(_lines) + "\n</user_long_term_memory>")
    # P2⑫ 云记忆自动画像：每轮读该用户自动画像常驻注入，让智能体据此调整语气/专业深度/关注重点。
    # 与 user_memory（零散显式记忆）互补，profile 是结构化、自动维护的「用户是谁」画像。异常降级为空。
    try:
        _prof = get_user_profile(u["id"]) if u else None
    except Exception:
        _prof = None
    if _prof:
        extra_sys += ("\n\n" if extra_sys else "") + (
            "<user_profile>\n以下是系统自动从你与该用户的历史对话中提取的「用户画像」（请据此调整语气、专业深度与关注重点）：\n"
            + _prof + "\n</user_profile>")
    # P1⑥ Plan 模式·执行阶段：把用户已确认的计划注入 system prompt，并指令用 Task* 跟踪进度
    if mode == "execute":
        _confirmed_plan = (payload.get("plan") or "").strip()
        if _confirmed_plan:
            extra_sys = (extra_sys + "\n\n" if extra_sys else "") + (
                "【已确认执行计划】\n" + _confirmed_plan +
                "\n\n请严格按上述已确认计划执行任务：使用 task_create 把每个步骤登记为子任务、"
                "用 task_update 维护状态（pending→in_progress→completed）、用 task_list 查看整体进度；"
                "完成全部步骤后再给出最终结论。不要偏离计划自行另起炉灶，也不要重复向用户确认已明确的步骤。"
            )
    if extra_sys:
        messages[0] = {"role": "system", "content": messages[0]["content"] + "\n\n" + extra_sys}
    if allowed_union:
        tool_defs = _narrow_tools_by_whitelist(tool_defs, tools_lib, allowed_union)

    # 技能分流：代码类走沙箱工具；方法论类已注入 system prompt，不再包装成工具
    code_defs = [d for d in skill_defs if d.get("skill_type") != "method"]
    tools = build_session_tools(ctx, tool_defs) + build_session_skill_tools(ctx, code_defs) + meta

    # 注：不再在路由层「全局剔除 create_skill」。
    # 旧做法依赖 method_text（关键词/语义匹配信号）判断是否禁用，既不可靠（多情况会漏匹配），
    # 又会误伤多技能、多轮任务（中途确实需要新建技能时被一刀切挡死）。
    # 重复创建的根治已下沉到工具层：_h_create_skill 在落库前基于「已有技能库」做状态级去重
    # （归一化名称相同 / 触发词重叠≥0.6 即视为重复，直接复用或自动安装，不创建）。
    # 因此 create_skill 在本任务中始终可用，且仅在真正需要时新建、不会重复。

    # 注册本任务（仅流式路径），供「停止」按钮通过 /api/agent/chat/abort 置位中止标志
    # 注意：task_id 已在上方匹配区统一定义（session_id 同源），此处不再重新生成，避免会话不一致。
    # 2026-08-19 隔离加固：存 user_id，abort 端点校验归属（避免任意用户中止别人的任务）
    AGENT_TASKS[task_id] = {"cancelled": False, "user_id": u["id"]}

    # 缺技能兜底·状态注入：前端在弹窗确认后用 task_phase 标识当前阶段，注入 system prompt
    # 让 LLM 直接进入对应执行模式，避免反复重新生成框架/标记，省 token 与时间。
    task_phase = (payload.get("task_phase") or "").strip()
    phase_hint = ""
    if task_phase == "framework_confirmed":
        phase_hint = ("\n\n【当前任务阶段】用户已在弹窗中确认思考框架（见最新用户消息中的"
                      "「[用户已确认思考框架]：...」）。请直接按确认的框架执行任务，"
                      "不要重新生成框架，不要再输出 [PROPOSE_METHOD] 标记。"
                      "任务完成后输出 [SUGGEST_SAVE_SKILL] 标记询问保存。")
    elif task_phase == "save_confirmed":
        phase_hint = ("\n\n【当前任务阶段】用户已同意把本次思考框架保存为技能（见最新用户消息中的"
                      "「是，请把本次思考框架保存为技能，display_name=XX，name=YY」）。"
                      "请立即调用 create_skill 工具：skill_type=method，"
                      "display_name/name 用用户消息中给出的值，instructions 填写用户在"
                      "「[用户已确认思考框架]」消息里给的框架全文。保存成功后用一句话告知用户即可，"
                      "不要再执行任务、不要再输出任何标记。")
    elif task_phase == "tool_confirmed":
        phase_hint = ("\n\n【当前任务阶段】用户已同意生成私有工具（见最新用户消息"
                      "「是，请创建该工具并立即执行任务」）。请立即调用 create_tool 工具："
                      "在 code 参数中写出 Python 代码（定义 def run(a): ... 且仅做纯数据处理，"
                      "禁止 os/sys/subprocess/socket/requests/eval/exec），name/display_name/description "
                      "按之前 [PROPOSE_TOOL] 标记里描述的工具用途给出合理值。保存成功后，"
                      "在剩余步骤中调用该工具执行任务，给出最终结论。不要再输出 [PROPOSE_TOOL] 标记。")
    if phase_hint:
        messages[0] = {"role": "system", "content": messages[0]["content"] + phase_hint}

    # P2⑮ 根因修复：把完整 messages 副本注入 ctx.messages，供工具入口在 LLM 未传
    # 必要参数时回灌聊天历史数据——典型场景：generate_ppt 收到 LLM 调它但 content="",
    # 此时工具自动从 ctx.messages 抽最近一条 assistant 文本作为 content 填入，
    # 解决"LLM 反思/重试只传 topic 不传 content → PPT 沦为通用空模板"问题。
    # 副本避免 run_agent 内部 append 时污染工具读到的 messages。
    try:
        ctx.messages = list(messages)
    except Exception:
        ctx.messages = None

    async def gen():
        yield _sse({"type": "start", "task_id": task_id, "tool_mode": tool_mode})
        # P1⑥ 规划模式：仅产出分步计划，不下发执行。用户在前端确认/编辑后，
        # 以 mode=execute + plan=已确认计划 续跑（新请求），进入真正执行。
        if mode == "plan":
            yield _sse({"type": "thinking", "text": "正在规划任务步骤…"})
            plan_text = await _produce_plan(client, name, params, question, catalog_text, method_text)
            yield _sse({"type": "plan_proposal", "text": plan_text})
            AGENT_TASKS.pop(task_id, None)
            return
        # ---- 已注入的方法论技能（method）：在步骤区显示「调用技能：XXX」----
        # method_text 已注入 system prompt 作为思考框架；这里把实际注入的技能名
        # 作为独立步骤渲染，让用户清晰看到「本次先使用了哪个技能」再进入工具调用。
        # 前端 plan/call(kind=skill)/result 已支持该展示，无需改动前端。
        for _ms in skill_defs:
            if _ms.get("skill_type") == "method":
                _ms_dn = _ms.get("display_name") or _ms.get("name") or "方法论技能"
                yield _sse({"type": "plan", "text": "调用技能：" + _ms_dn})
                yield _sse({"type": "call", "kind": "skill", "tool": _ms_dn})
                yield _sse({"type": "result"})
        # ---- 未安装技能检测（技能优先层）----
        # 用户即便解除了技能安装，只要广场里存在公开且关键词命中的技能，
        # 就主动提示安装，而不是让 LLM 因检索不到而走 create_skill 兜底。
        # 属于技能优先层（优先级最高），与「工具全量/关键词」开关无关，两种模式都生效。
        # 仅当「已安装技能均未命中」时才提示未安装技能，避免已安装技能能覆盖时误弹窗。
        # 【防误触发】纯闲聊 / 仅含问候语的消息不触发安装推荐（避免日常对话误弹窗）。
        if not _is_greeting_or_chitchat(question):
            _all_sk = list_skills(for_user_id=u["id"])  # 已发布+本人私有，含 installed 标注
            _installed_hit = any(_has_lexical_hit(question, [s]) for s in _all_sk if s.get("installed"))
            if not _installed_hit:
                _uninst = None
                for _s in _all_sk:
                    if (not _s.get("installed")
                            and _s.get("status") == "approved" and _s.get("scope") == "public"
                            and _has_lexical_hit(question, [_s])):
                        _uninst = _s
                        break
                if _uninst:
                    yield _sse({"type": "propose_skill_install",
                                "id": _uninst["id"],
                                "name": _uninst["name"],
                                "display_name": _uninst.get("display_name") or _uninst["name"],
                                "description": _uninst.get("description", ""),
                                "trigger_words": _uninst.get("trigger_words", "")})
                    # 不 return——用户选"暂不使用"时继续走 LLM 兜底（create_skill / 自身能力），
                    # 任务完成后前端会弹出"保存为新技能"确认框（同广场无技能时的兜底方式）
        # 回显本次实际生效的模型，便于用户确认下拉切换是否真正生效
        yield _sse({"type": "model_used",
                    "name": active.get("name") or "",
                    "model_name": name,
                    "picked": bool(sel_id)})
        if bool(payload.get("web")):
            if web_hits:
                yield _sse({"type": "web_search", "status": "ok",
                            "count": len(web_hits), "engine": web_engine,
                            "sources": [{"t": h.get("title", ""), "u": h.get("url", "")}
                                        for h in web_hits]})
            else:
                yield _sse({"type": "web_search", "status": "empty", "count": 0})
        if constraint_text:
            yield _sse({"type": "skill_constraint", "text": constraint_text,
                        "allowed_tools": sorted(allowed_union)})
        # 第一遍：仅 yield tool_injected 通知事件（不执行 agent 循环）
        for t in selected:
            is_method = (t.get("skill_type") == "method")
            kd = "method" if is_method else (t.get("kind") or "tool")
            sd = (not is_method) and (t.get("kind") == "skill")
            yield _sse({"type": "tool_injected", "tool": t["name"],
                        "description": t.get("description", ""),
                        "kind": kd,
                        "sandboxed": sd})

        # 第二遍：只调用一次 run_agent()，把所有选中工具一次性传入 ReAct 循环
        # （之前的 bug：run_agent 被放在 for t in selected 内部循环里，
        #   导致 selected 有 N 项就跑 N 次完整 ReAct 循环，造成事件重复 + 极慢 + 流不关闭）
        yield _sse({"type": "thinking", "text": "正在调用大模型分析任务…"})
        _is_cancel = lambda: AGENT_TASKS.get(task_id, {}).get("cancelled", False)
        _max_steps = int((active.get("max_steps") or 8) or 8)
        if _max_steps < 1:
            _max_steps = 8
        try:
            async for ev in run_agent_with_retry(client, name, messages, tools, params, max_steps=_max_steps,
                                      on_tool_call=_on_agent_call, cancel_check=_is_cancel, ctx=ctx,
                                      caps=caps, original_question=question, max_attempts=2):
                # 客户端断开连接则立即退出，避免服务端空转
                if await request.is_disconnected():
                    break
                if ev.get("type") == "aborted":
                    yield _sse({"type": "aborted"})
                    break
                # 文件型工具会把产物路径写进返回文本，检测后对每个真实产物向前端推送下载卡片
                if ev.get("type") == "result":
                    for fp in _detect_generated_files(ev.get("text", "")):
                        yield _sse({"type": "file", "path": fp,
                                    "name": os.path.basename(fp),
                                    "size": os.path.getsize(fp),
                                    "ext": _ext_of(fp),
                                    "tool": ev.get("tool", "")})
                # 缺技能/缺工具兜底：解析 final 事件里的 [PROPOSE_METHOD] / [SUGGEST_SAVE_SKILL] / [PROPOSE_TOOL] 标记，
                # 从 text 剔除后下发 final，再额外 yield 专用事件触发前端弹窗
                if ev.get("type") == "final":
                    clean, propose, suggest, propose_tool = _parse_method_markers(ev.get("text", ""))
                    if clean != ev.get("text"):
                        ev = dict(ev, text=clean)
                    yield _sse(ev)
                    # P2⑪ 对话消息落库：助手本轮终答写入 conversations 表
                    try:
                        db.save_conversation_message(session_id, u["id"], "assistant", clean)
                    except Exception as _e2:
                        print(f"[P2⑪] 助手回复落库跳过: {_e2}")
                    # P2⑫ 云记忆自动画像：终答落库后后台抽取（不阻塞 SSE 流）。
                    # 失败静默——画像只是增强，绝不影响主任务与回复。
                    try:
                        _pt = asyncio.create_task(
                            _run_profile_extraction(client, name, params, u["id"], question, clean))
                        _PROFILE_TASKS.add(_pt)
                        _pt.add_done_callback(_PROFILE_TASKS.discard)
                    except Exception as _e3:
                        print(f"[P2⑫] profile task skip: {_e3}")
                    # 终答里若提到真实存在的生成文件，也补一张下载卡片（前端按路径去重，避免与 result 重复）
                    for fp in _detect_generated_files(clean):
                        yield _sse({"type": "file", "path": fp,
                                    "name": os.path.basename(fp),
                                    "size": os.path.getsize(fp),
                                    "ext": _ext_of(fp),
                                    "tool": ev.get("tool", "")})
                    if propose:
                        yield _sse({"type": "propose_method", "framework": propose})
                    if suggest:
                        dn, nm = suggest
                        yield _sse({"type": "suggest_save_skill",
                                    "display_name": dn, "name": nm})
                    if propose_tool:
                        yield _sse({"type": "propose_tool", "description": propose_tool})
                else:
                    yield _sse(ev)
            for sk in getattr(ctx, "created_skills", []):
                yield _sse({"type": "skill_created", "name": sk["name"], "id": sk["id"]})
            for tl in getattr(ctx, "created_tools", []):
                yield _sse({"type": "tool_created", "name": tl["name"], "id": tl["id"]})
        except asyncio.CancelledError:
            # 客户端中途断开（刷新页面 / 中止读取），正常退出
            pass
        finally:
            AGENT_TASKS.pop(task_id, None)
        add_log("agent_chat", detail=question[:50], user=u, ip=client_ip(request))

    return StreamingResponse(gen(), media_type="text/event-stream")


@router.post("/api/agent/chat/abort")
async def agent_chat_abort(request: Request, payload: dict):
    """中止正在运行的智能体任务（前端「停止」按钮调用）。

    通过 task_id 找到对应任务并置位 cancelled 标志；gen()/run_agent() 在下一个
    决策点（轮次 / 工具调用）感知后优雅退出，避免服务端继续空转消耗 LLM 配额。
    **2026-08-19 隔离加固**：仅当 task_id 归属当前用户时才能中止；否则返回 403。
    """
    u = require_perm("agent", request)
    tid = (payload or {}).get("task_id") or ""
    rec = AGENT_TASKS.get(tid)
    if rec is not None:
        if rec.get("user_id") != u["id"]:
            # 隔离：不是自己的任务不能中止
            return JSONResponse(status_code=403, content={"ok": False, "error": "无权中止该任务（归属不匹配）"})
        rec["cancelled"] = True
        return JSONResponse(content={"ok": True, "cancelled": True})
    return JSONResponse(content={"ok": True, "cancelled": False,
                                 "note": "任务不存在或已完成"})


@router.get("/api/agent/download")
async def agent_download(path: str, request: Request):
    """下载智能体生成的文件。白名单目录 + 用户归属校验（防路径穿越 + 跨用户越权）。"""
    u = require_perm("agent", request)
    if not path:
        return JSONResponse(status_code=400, content={"error": "缺少 path"})
    real = os.path.realpath(path)
    # 用户归属校验：普通用户只能下载自己 <user_id>/ 子目录下的文件（管理员放行；旧平铺无主文件拒绝）
    if not _file_belongs_to(u, real):
        return JSONResponse(status_code=403, content={"error": "无权下载该文件（仅能下载本人生成的文件）"})
    roots = [os.path.realpath(tempfile.gettempdir())]
    extra = (os.environ.get("AGENT_DOWNLOAD_ROOTS") or "").strip()
    if extra:
        roots += [os.path.realpath(r.strip()) for r in extra.split(",") if r.strip()]
    # 沙箱产物持久目录（与 sandbox.py 的 ARTIFACT_ROOT 保持一致）
    _art_root = os.environ.get("SKILL_ARTIFACT_ROOT") or (
        "/app/data/artifacts" if os.path.isdir("/app/data") else tempfile.gettempdir())
    roots.append(os.path.realpath(_art_root))
    if not any(real == rp or real.startswith(rp + os.sep) for rp in roots):
        return JSONResponse(status_code=403, content={"error": "路径不在允许下载范围内"})
    if not os.path.isfile(real):
        return JSONResponse(status_code=404, content={"error": "文件不存在或已被清理"})
    return FileResponse(real, filename=os.path.basename(real),
                        media_type="application/octet-stream")


@router.get("/api/agent/chat/models")
async def agent_chat_models(request: Request):
    """当前用户可用的模型：role=chat（主推理，下拉默认）+ role=vision（多模态生成，可选）。
    is_active=1 为当前全局默认。"""
    require_perm("agent", request)
    c = get_conn()
    rows = c.execute(
        "SELECT id, name, model_name, role, is_active FROM models "
        "WHERE role IN ('chat','vision') AND enabled=1 ORDER BY role, is_active DESC, id"
    ).fetchall()
    c.close()
    models = [dict(r) for r in rows]
    # 库内无 chat 记录时，回落到环境变量种子配置（get_active 的兜底路径），避免下拉空白
    if not any(m.get("role") == "chat" for m in models):
        act = get_active("chat") or {}
        if act.get("model_name"):
            models.append({"id": 0, "name": act.get("name") or "环境变量配置",
                           "model_name": act.get("model_name"), "role": "chat", "is_active": 1})
    return JSONResponse(content={"models": models})


@router.get("/api/agent/memories")
async def agent_list_memories(request: Request):
    """查看当前用户的跨会话长期记忆列表（前端「我的记忆」面板用）。"""
    u = require_perm("agent", request)
    try:
        mems = get_user_memory(u["id"]) if u else []
    except Exception as e:
        return JSONResponse(content={"ok": False, "error": str(e), "memories": []}, status_code=200)
    return JSONResponse(content={"ok": True, "memories": mems})


@router.delete("/api/agent/memories")
async def agent_delete_memory(request: Request, payload: dict = None):
    """删除当前用户的一条长期记忆（按 mem_id 或 mem_key）。"""
    u = require_perm("agent", request)
    p = payload or {}
    mem_id = p.get("mem_id") or p.get("id")
    mem_key = p.get("mem_key")
    try:
        delete_user_memory(u["id"], mem_id=mem_id, mem_key=mem_key)
    except Exception as e:
        return JSONResponse(content={"ok": False, "error": str(e)}, status_code=200)
    return JSONResponse(content={"ok": True})


@router.get("/api/agent/profile")
async def agent_get_profile(request: Request):
    """查看当前用户的自动画像（前端「我的画像」面板用）。"""
    u = require_perm("agent", request)
    try:
        _conn = get_conn()
        try:
            _row = _conn.execute(
                "SELECT profile, updated_at FROM user_profiles WHERE user_id=?", (u["id"],)).fetchone()
        finally:
            _conn.close()
        prof = _row["profile"] if _row else None
        updated_at = _row["updated_at"] if _row else None
    except Exception as e:
        return JSONResponse(content={"ok": False, "error": str(e), "profile": None, "updated_at": None}, status_code=200)
    return JSONResponse(content={"ok": True, "profile": prof, "updated_at": updated_at})


@router.delete("/api/agent/profile")
async def agent_clear_profile(request: Request, payload: dict = None):
    """清空当前用户的自动画像（重新对话后会再次自动累积）。"""
    u = require_perm("agent", request)
    try:
        save_user_profile(u["id"], "")
    except Exception as e:
        return JSONResponse(content={"ok": False, "error": str(e)}, status_code=200)
    return JSONResponse(content={"ok": True})


@router.get("/api/agent/tasks")
async def agent_list_tasks(request: Request, session_id: str = None):
    """查看当前用户的任务列表（前端任务进度面板用）。可传 session_id 仅看本次连续对话。"""
    u = require_perm("agent", request)
    try:
        rows = list_tasks(u["id"], session_id=session_id)
    except Exception as e:
        return JSONResponse(content={"ok": False, "error": str(e), "tasks": []}, status_code=200)
    return JSONResponse(content={"ok": True, "tasks": rows})


# ────────────────────────────────────────────────────────────────────────
# P1④ · MCP 连接器：查看 / 手动刷新
# ────────────────────────────────────────────────────────────────────────
@router.get("/api/agent/mcp/tools")
async def agent_list_mcp(request: Request):
    """查看已同步进能力目录的 MCP 工具（来自 ~/.workbuddy/mcp.json）。"""
    u = require_perm("agent", request)
    if mcp_client is None:
        return JSONResponse(content={"ok": False, "error": "MCP 模块未加载", "tools": []}, status_code=200)
    try:
        rows = [r for r in list_tools(u["id"]) if (r.get("backend_type") == "mcp")]
    except Exception as e:
        return JSONResponse(content={"ok": False, "error": str(e), "tools": []}, status_code=200)
    return JSONResponse(content={"ok": True, "tools": rows, "configured": mcp_client.load_mcp_config()})


@router.post("/api/agent/mcp/refresh")
async def agent_refresh_mcp(request: Request):
    """手动重新发现并同步 MCP 工具到能力目录。"""
    u = require_perm("agent", request)
    if mcp_client is None:
        return JSONResponse(content={"ok": False, "error": "MCP 模块未加载"}, status_code=200)
    try:
        n = mcp_client.sync_mcp_tools()
    except Exception as e:
        return JSONResponse(content={"ok": False, "error": str(e)}, status_code=200)
    return JSONResponse(content={"ok": True, "synced": n})


# ────────────────────────────────────────────────────────────────────────
# P2⑩ · 定时自动化：管理端点（用户级）
# ────────────────────────────────────────────────────────────────────────
@router.get("/api/agent/automations")
async def agent_list_automations(request: Request):
    """列出当前用户的全部自动化任务，并附带最近一次执行摘要。"""
    u = require_perm("agent", request)
    try:
        rows = list_automations(owner_id=u["id"])
        for r in rows:
            runs = list_automation_runs(r["id"], limit=1)
            r["last_run"] = runs[0] if runs else None
    except Exception as e:
        return JSONResponse(content={"ok": False, "error": str(e), "automations": []}, status_code=200)
    return JSONResponse(content={"ok": True, "automations": rows})


@router.post("/api/agent/automations")
async def agent_create_automation(request: Request, payload: dict = None):
    """创建自动化任务。body: name,prompt,schedule_type(recurring|once),rrule?,scheduled_at?,valid_from?,valid_until?,status?"""
    u = require_perm("agent", request)
    p = payload or {}
    name = (p.get("name") or "").strip()
    prompt = (p.get("prompt") or "").strip()
    if not name or not prompt:
        return JSONResponse(content={"ok": False, "error": "name 与 prompt 必填"}, status_code=200)
    st = (p.get("schedule_type") or "recurring")
    if st not in ("recurring", "once"):
        st = "recurring"
    status = (p.get("status") or "ACTIVE").upper()
    if status not in ("ACTIVE", "PAUSED"):
        status = "ACTIVE"
    try:
        aid = create_automation(
            owner_id=u["id"], name=name, prompt=prompt, schedule_type=st,
            rrule=(p.get("rrule") or ""), scheduled_at=p.get("scheduled_at") or None,
            valid_from=p.get("valid_from") or None, valid_until=p.get("valid_until") or None,
            status=status, notify_email=(p.get("notify_email") or "").strip())
    except Exception as e:
        return JSONResponse(content={"ok": False, "error": str(e)}, status_code=200)
    return JSONResponse(content={"ok": True, "id": aid})


@router.patch("/api/agent/automations/{aid}")
async def agent_update_automation(aid: int, request: Request, payload: dict = None):
    """更新自动化（含激活/暂停 status 切换、改 prompt/rrule 等），需所有权校验。"""
    u = require_perm("agent", request)
    p = payload or {}
    if not get_automation(aid, owner_id=u["id"]):
        return JSONResponse(content={"ok": False, "error": "任务不存在或无权访问"}, status_code=200)
    try:
        update_automation(aid,
            name=p.get("name"), prompt=p.get("prompt"),
            schedule_type=p.get("schedule_type"), rrule=p.get("rrule"),
            scheduled_at=p.get("scheduled_at"), status=p.get("status"),
            valid_from=p.get("valid_from"), valid_until=p.get("valid_until"),
            notify_email=p.get("notify_email"))
    except Exception as e:
        return JSONResponse(content={"ok": False, "error": str(e)}, status_code=200)
    return JSONResponse(content={"ok": True})


@router.delete("/api/agent/automations/{aid}")
async def agent_delete_automation(aid: int, request: Request):
    """删除自动化任务及其执行历史，需所有权校验。"""
    u = require_perm("agent", request)
    if not get_automation(aid, owner_id=u["id"]):
        return JSONResponse(content={"ok": False, "error": "任务不存在或无权访问"}, status_code=200)
    try:
        delete_automation(aid)
    except Exception as e:
        return JSONResponse(content={"ok": False, "error": str(e)}, status_code=200)
    return JSONResponse(content={"ok": True})


@router.post("/api/agent/automations/{aid}/run")
async def agent_run_automation(aid: int, request: Request):
    """手动立即触发一次自动化执行（与定时调度互不影响）。"""
    u = require_perm("agent", request)
    if not get_automation(aid, owner_id=u["id"]):
        return JSONResponse(content={"ok": False, "error": "任务不存在或无权访问"}, status_code=200)
    try:
        asyncio.create_task(automation_runner.run_automation(aid))
    except Exception as e:
        return JSONResponse(content={"ok": False, "error": str(e)}, status_code=200)
    return JSONResponse(content={"ok": True, "message": "已触发执行"})


# ────────────────────────────────────────────────────────────────────────
# P2⑪ 档B：历史对话（会话窗口）管理
# 列表 / 重命名 / 级联删除 / 续聊取历史。所有权均按当前用户校验。
# ────────────────────────────────────────────────────────────────────────

@router.get("/api/agent/conversations")
async def agent_list_conversations(request: Request):
    """列出当前用户的所有历史对话窗口（含消息数、最后活动时间）。"""
    try:
        u = require_perm("agent", request)
        uid = u["id"]
        # 诊断日志：定位为什么 sessions 返回空
        conn = db.get_conn()
        total = conn.execute("SELECT COUNT(*) FROM conversation_sessions").fetchone()[0]
        mine = conn.execute("SELECT COUNT(*) FROM conversation_sessions WHERE user_id=?", [uid]).fetchone()[0]
        sample = conn.execute("SELECT session_id, user_id, title, deleted_at FROM conversation_sessions LIMIT 3").fetchall()
        print(f"[P2⑪ DIAG] uid={uid}, total_sessions={total}, mine={mine}, sample={[dict(r) for r in sample]}")
        conn.close()
        rows = db.list_sessions(uid)
        print(f"[P2⑪ DIAG] list_sessions returned {len(rows)} rows")
        return JSONResponse(content={"ok": True, "sessions": rows})
    except HTTPException:
        raise
    except Exception as e:
        return JSONResponse(content={"ok": False, "error": str(e)}, status_code=500)


@router.patch("/api/agent/conversations/{sid}")
async def agent_rename_conversation(sid: str, request: Request):
    """重命名某个历史对话窗口。**用户隔离**：仅当 sid 归属当前用户时生效。"""
    try:
        u = require_perm("agent", request)
        try:
            body = await request.json()
        except Exception:
            body = {}
        title = (body.get("title") or "").strip()
        if not title:
            return JSONResponse(content={"ok": False, "error": "标题不能为空"}, status_code=200)
        # 2026-08-19 隔离加固：传 user_id 校验归属，避免任意用户改别人的会话标题
        ok = db.update_session_title(u["id"], sid, title)
        if not ok:
            return JSONResponse(content={"ok": False, "error": "会话不存在或无权访问"}, status_code=200)
        return JSONResponse(content={"ok": True})
    except HTTPException:
        raise
    except Exception as e:
        return JSONResponse(content={"ok": False, "error": str(e)}, status_code=500)


@router.delete("/api/agent/conversations/{sid}")
async def agent_delete_conversation(sid: str, request: Request):
    """级联删除某个历史对话窗口及其全部消息（删除=清历史）。"""
    try:
        u = require_perm("agent", request)
        ok = db.delete_session(u["id"], sid)
        if not ok:
            return JSONResponse(content={"ok": False, "error": "会话不存在或无权访问"}, status_code=200)
        return JSONResponse(content={"ok": True})
    except HTTPException:
        raise
    except Exception as e:
        return JSONResponse(content={"ok": False, "error": str(e)}, status_code=500)


@router.get("/api/agent/conversations/{sid}/messages")
async def agent_get_conversation_messages(sid: str, request: Request):
    """取某历史对话窗口的全部消息（续聊时前端恢复上下文用）。"""
    try:
        u = require_perm("agent", request)
        msgs = db.get_session_messages(u["id"], sid)
        return JSONResponse(content={"ok": True, "messages": msgs, "session_id": sid})
    except HTTPException:
        raise
    except Exception as e:
        return JSONResponse(content={"ok": False, "error": str(e)}, status_code=500)


# ────────────────────────────────────────────────────────────────────────
# P2 · 历史窗口压缩（compaction）
# 设计：同 session 长对话中，窗口内（最近 keep 条）逐字保留，窗口外（更早的）
# 压缩为一条 system 摘要注入，避免早期任务背景完全丢失（原 history[-10:] 硬截断直接丢弃）。
# 摘要采用规则启发式（零额外 LLM 调用、完美兼容任意模型、绝不因摘要失败而终止任务）；
# 若日后需要更高质量，可在 _summarize_older 接入 LLM（预留扩展点）。
# 注意：压缩只作用于「注入给 LLM 的历史」，不影响 P1 的 session_active_caps 持久化。
# ────────────────────────────────────────────────────────────────────────

def _summarize_older(older, max_chars=1400):
    """把窗口外的历史消息压缩为结构化摘要文本。任何异常返回空串（调用方降级处理）。"""
    try:
        user_reqs, artifacts, actions = [], [], []
        for m in older:
            c = (m.get("content") or "").strip()
            if not c:
                continue
            if m.get("role") == "user":
                snippet = re.sub(r"\s+", " ", c)[:240]
                if snippet:
                    user_reqs.append(snippet)
            else:  # assistant
                for p in re.findall(r'(?:已生成|已保存|保存至|保存到|存放于|输出到|输出至|路径)[:：]\s*(\S+)', c):
                    if p not in artifacts:
                        artifacts.append(p)
                for t in re.findall(r'调用技能[：:]\s*([^\n\[】]+)', c):
                    t = t.strip().rstrip('】').strip()
                    if t and t not in actions:
                        actions.append(t)
                if '✅ 已创建成功' in c or '工具已创建成功' in c or '技能已创建' in c:
                    if '（曾创建工具/技能）' not in actions:
                        actions.append('（曾创建工具/技能）')
        parts = []
        if user_reqs:
            parts.append("用户早期诉求：\n" + "\n".join(f"  - {r}" for r in user_reqs[:10]))
        if actions:
            parts.append("已执行动作：" + "；".join(actions[:12]))
        if artifacts:
            parts.append("已完成产物（路径）：\n" + "\n".join(f"  - {a}" for a in artifacts[:20]))
        text = "\n".join(parts)
        if not text.strip():
            return ""
        if len(text) > max_chars:
            text = text[:max_chars] + "\n…(早期摘要已截断)"
        return ("[对话早期摘要] 以下为本次长对话中超出上下文窗口的早期内容压缩摘要，"
                "用于保留任务背景，细节可能不完整；如需精确信息，可基于磁盘产物（如上述路径文件）重新读取：\n" + text)
    except Exception:
        return ""


def _compact_history(history, keep=12, max_summary_chars=1400):
    """窗口内逐字 + 窗口外摘要。返回可直接 append 的消息片段列表。

    - 最近 keep 条 user/assistant 消息：原样保留（逐字）
    - 更早的消息：经 _summarize_older 压缩为一条 system 摘要
    - 任何异常均降级为「最近 keep 条原样」（等价于原 history[-keep:] 行为），绝不抛错终止任务
    """
    try:
        if not isinstance(history, list):
            return []
        cleaned = [m for m in history
                   if isinstance(m, dict) and m.get("role") in ("user", "assistant") and m.get("content")]
        if len(cleaned) <= keep:
            return cleaned
        window = cleaned[-keep:]
        older = cleaned[:-keep]
        summary = _summarize_older(older, max_summary_chars)
        if not summary:
            return window
        return [{"role": "system", "content": summary}] + window
    except Exception:
        # 降级：原行为（最近 keep 条）
        try:
            if isinstance(history, list):
                return [m for m in history
                        if isinstance(m, dict) and m.get("role") in ("user", "assistant") and m.get("content")][-keep:]
        except Exception:
            pass
        return []


def _render_capability_catalog(tools_lib, skills_lib):
    """生成「常驻能力目录」文本块（全量轻量，每轮注入 system prompt）。

    只列 name + 中文名 + 一句话描述 + 触发词，绝不预载完整 schema/指令（控 token）。
    让 LLM 每轮都看到「完整可用能力菜单」，基于其原生推理决定调用哪个能力——
    这是「匹配=LLM 推理式」的载体（§1.1）：匹配信号永远「增加」候选，不靠关键词剥夺可用性。
    """
    if not (tools_lib or skills_lib):
        return ""
    lines = ["【你可用的能力清单（常驻·每轮可见，完整版）】"]
    if tools_lib:
        lines.append("一、工具（按需在 tool_call 中调用）：")
        for t in tools_lib:
            dn = t.get("display_name") or t.get("name") or ""
            tw = (t.get("trigger_words") or "").replace(",", " ").strip()
            desc = (t.get("description") or "").strip()
            lines.append(f"- {t['name']}（{dn}）：{desc}　触发词：{tw}")
    if skills_lib:
        lines.append("二、技能（method=思考框架将注入提示；code=代码技能可调用）：")
        for s in skills_lib:
            st = s.get("skill_type") or "method"
            dn = s.get("display_name") or s.get("name") or ""
            tw = (s.get("trigger_words") or "").replace(",", " ").strip()
            desc = (s.get("description") or "").strip()
            lines.append(f"- {s['name']}（{dn}，{st}）：{desc}　触发词：{tw}")
    lines.append("（以上为完整能力清单，请基于当前任务自行判断调用哪个；无需再调用 "
                 "list_available_skills / list_available_tools 来发现能力，需要时可直接 tool_call。）")
    return "\n".join(lines)


def _on_agent_call(tool_name):
    """工具/技能执行计数（按类型分流到对应表）。"""
    inc_tool_calls(tool_name)
    inc_skill_calls(tool_name)


# ────────────────────────────────────────────────────────────────────────────
# P2⑫ 云记忆自动画像：对话终答后由 LLM 自动抽取并写入 user_profiles，每轮注入 system prompt。
# 后台异步执行，不阻塞 SSE 流；任何异常静默——画像只是增强，绝不影响主任务与回复。
# ────────────────────────────────────────────────────────────────────────────
_PROFILE_TASKS = set()


def _extract_profile_sync(client, name, params, user_id, question, answer):
    """同步抽取：读旧画像 + 本轮对话 → LLM 合并更新 → 写回。异常静默。"""
    old = ""
    try:
        old = get_user_profile(user_id) or ""
    except Exception:
        old = ""
    sys_p = (
        "你是一个用户画像抽取助手。根据「用户与助手的本轮对话」，维护该用户的长期画像。\n"
        "画像用简洁中文 Markdown 结构化描述，分以下小节（无内容的小节可省略）：\n"
        "## 角色与身份\n## 业务领域\n## 偏好与习惯\n## 约束与禁忌\n## 项目背景\n"
        "规则：\n"
        "1) 仅写入从对话中可明确推断的稳定事实，严禁编造；\n"
        "2) 若本轮未透露新信息，则保持原画像不变（直接原样输出旧画像）；\n"
        "3) 合并去重，同一主题只保留一条最新结论，不要罗列历史；\n"
        "4) 只写稳定画像，不要写具体对话内容/问答原文；\n"
        "5) 直接输出更新后的【完整画像】Markdown，不要再附加任何解释。"
    )
    user_p = (
        f"【现有画像】\n{(old if old else '(无)')}\n\n"
        f"【本轮对话】\n用户：{question}\n助手：{answer}\n\n"
        "请输出更新后的完整画像："
    )
    try:
        resp = client.chat.completions.create(
            model=name,
            messages=[{"role": "system", "content": sys_p},
                      {"role": "user", "content": user_p}],
            temperature=0.1,
            max_tokens=800)
        new_prof = (resp.choices[0].message.content or "").strip()
    except Exception as e:
        print(f"[P2⑫] LLM extract error: {e}")
        return
    if new_prof and new_prof != old.strip():
        try:
            save_user_profile(user_id, new_prof)
        except Exception as e:
            print(f"[P2⑫] save profile error: {e}")


async def _run_profile_extraction(client, name, params, user_id, question, answer):
    """后台任务：在线程池跑同步 LLM 抽取，避免阻塞事件循环 / SSE 流。"""
    try:
        await asyncio.to_thread(_extract_profile_sync, client, name, params, user_id, question, answer)
    except Exception as e:
        print(f"[P2⑫] profile extract skip: {e}")

