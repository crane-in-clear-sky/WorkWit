"""Word 生成：将文本内容（如诗歌、文章、报告）生成 Word（.docx）文件，支持丰富排版。"""
import os
import re
import tempfile
import uuid

from builtin_tools._shared import _user_dir, _safe_id

META = {
    "name": "generate_word", "display_name": "Word 生成", "category": "generation",
    "description": "将文本内容（如诗歌、文章、报告）生成 Word（.docx）文件并支持下载；当用户要求「输出为 Word / 导出成文档 / 保存为 .docx」时使用。",
    "params": {"type": "object",
               "properties": {"content": {"type": "string", "description": "要写入 Word 文档的正文内容（支持多段落，用换行分隔）"},
                              "title": {"type": "string", "description": "文档标题（可选，作为一级标题）"}},
               "required": ["content"]},
    "backend_type": "builtin", "handler": "generate_word",
    "trigger_words": "Word,word,docx,文档,.doc,导出word,生成word,保存为word,输出为word,word文档,word文件",
}


def run(ctx, content, title=""):
    """将文本内容生成精美排版的 Word（.docx）文件。

    渲染引擎支持：
    - Markdown 表格（| 列 | 列 |）→ 真实 Word 表格（带色块、边框）
    - 标题层级（# / ## / ###）→ 样式化标题（带颜色、左边框装饰）
    - 有序列表 / 无序列表（● / - / * / 1.）→ 彩色列表
    - 内联格式（**粗体** / *斜体* / `代码`）
    - 状态标记（✅绿 ⚠️橙 ✗红 🔵蓝）→ 保留彩色
    - 分隔线（--- / ***）→ 水平线
    - 整体文档样式（页边距、字体、段落间距）
    """
    _uid = _safe_id(ctx.user.get("id") if getattr(ctx, "user", None) else None)
    _art_root = _user_dir(_uid, getattr(ctx, "session_id", None), "artifacts")
    try:
        os.makedirs(_art_root, exist_ok=True)
    except Exception:
        _art_root = tempfile.gettempdir()
    _safe = re.sub(r"\W+", "_", (title or "文档")[:30]) or "文档"
    out = os.path.join(_art_root, f"{_safe}_{uuid.uuid4().hex[:8]}.docx")

    try:
        import docx
        from docx.shared import Pt, Cm, RGBColor, Emu
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
        from docx.oxml.ns import qn, nsdecls
        from docx.oxml import parse_xml

        doc = docx.Document()

        # ====== 全局文档样式 ======
        section = doc.sections[0]
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

        style = doc.styles["Normal"]
        font = style.font
        font.name = "微软雅黑"
        font.size = Pt(11)
        style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

        C_GREEN = RGBColor(0x16, 0xA3, 0x4A)
        C_ORANGE = RGBColor(0xF5, 0x9E, 0x0B)
        C_RED = RGBColor(0xEF, 0x44, 0x44)
        C_BLUE = RGBColor(0x3B, 0x82, 0xF6)
        C_PURPLE = RGBColor(0x8B, 0x5C, 0xF6)
        C_GRAY = RGBColor(0x6B, 0x72, 0x80)
        C_BG_LIGHT = RGBColor(0xF8, 0xFA, 0xFC)
        C_BORDER = RGBColor(0xE2, 0xE8, 0xF0)

        def _set_cell_shading(cell, color_hex):
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
            cell._tc.get_or_add_tcPr().append(shading)

        def _add_styled_run(para, text, bold=False, italic=False, color=None,
                            font_name="微软雅黑", size=Pt(11)):
            r = para.add_run(text)
            r.bold = bold
            r.italic = italic
            if color:
                r.font.color.rgb = color
            r.font.name = font_name
            r.font.size = size
            r.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
            return r

        def _add_horizontal_rule(doc):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                '  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="CBD5E1"/>'
                '</w:pBdr>'
            )
            pPr.append(pBdr)

        def _add_blockquote(doc, text):
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                '  <w:left w:val="single" w:sz="18" w:space="8" w:color="CBD5E1"/>'
                '</w:pBdr>'
            )
            pPr.append(pBdr)
            shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:fill="F8FAFC"/>')
            pPr.append(shd)
            p.paragraph_format.left_indent = Cm(0.6)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            first = True
            for ln in (text or "").split("\n"):
                if not first:
                    p.add_run().add_break()
                first = False
                r = p.add_run(ln)
                r.italic = True
                r.font.name = "微软雅黑"
                r.font.size = Pt(10.5)
                r.font.color.rgb = C_GRAY
                r.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
            return p

        def _add_code_block(doc, code_text):
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:fill="F1F5F9"/>')
            pPr.append(shd)
            p.paragraph_format.left_indent = Cm(0.4)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            first = True
            for ln in (code_text or "").split("\n"):
                if not first:
                    p.add_run().add_break()
                first = False
                r = p.add_run(ln)
                r.font.name = "Consolas"
                r.font.size = Pt(10)
                r.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
                r.element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
                r.element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
            return p

        def _parse_table_row(line):
            stripped = line.strip()
            if not stripped.startswith("|") or not stripped.endswith("|"):
                return None
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            if all(re.match(r"^[\s\-:]+$", c) for c in cells):
                return None
            return cells

        def _render_table(doc, rows_data):
            if not rows_data:
                return
            ncols = max(len(r) for r in rows_data)
            table = doc.add_table(rows=len(rows_data), cols=ncols)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            for ri, row_data in enumerate(rows_data):
                row = table.rows[ri]
                for ci, cell_text in enumerate(row_data):
                    if ci >= len(row.cells):
                        break
                    cell = row.cells[ci]
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    cell.text = ""
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

                    is_header = (ri == 0)
                    if is_header:
                        _set_cell_shading(cell, "F1F5F9")
                        _add_styled_run(p, cell_text, bold=True,
                                       color=C_GRAY, size=Pt(10))
                    else:
                        text_color = None
                        clean_text = cell_text
                        if "✅" in cell_text:
                            text_color = C_GREEN
                        elif "⚠️" in cell_text or "⚠" in cell_text:
                            text_color = C_ORANGE
                        elif "✗" in cell_text or "❌" in cell_text:
                            text_color = C_RED
                        _add_styled_run(p, clean_text, color=text_color or None,
                                       size=Pt(10))

            for row in table.rows:
                for cell in row.cells:
                    tcPr = cell._tc.get_or_add_tcPr()
                    tcBorders = parse_xml(
                        f'<w:tcBorders {nsdecls("w")}>'
                        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>'
                        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>'
                        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>'
                        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>'
                        f'</w:tcBorders>'
                    )
                    tcPr.append(tcBorders)

            doc.add_paragraph()

        # ====== 主渲染循环 ======
        raw = content or ""
        if not raw.strip():
            doc.add_paragraph("（无内容）")
        else:
            if title:
                h = doc.add_heading(title, level=0)
                for run_ in h.runs:
                    run_.font.color.rgb = C_PURPLE
                    run_.font.name = "微软雅黑"
                    run_.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

            lines = raw.split("\n")
            i = 0
            pending_table_rows = []

            def _flush_table():
                nonlocal pending_table_rows
                if pending_table_rows:
                    _render_table(doc, pending_table_rows)
                    pending_table_rows = []

            while i < len(lines):
                line = lines[i].rstrip()

                if not line.strip():
                    _flush_table()
                    i += 1
                    continue

                if re.match(r"^[\-\*]{3,}\s*$", line.strip()):
                    _flush_table()
                    _add_horizontal_rule(doc)
                    i += 1
                    continue

                if line.strip().startswith("```"):
                    _flush_table()
                    code_lines = []
                    i += 1
                    while i < len(lines):
                        cl = lines[i].rstrip()
                        if cl.strip().startswith("```"):
                            break
                        code_lines.append(cl)
                        i += 1
                    i += 1
                    _add_code_block(doc, "\n".join(code_lines))
                    continue

                if line.strip().startswith(">"):
                    _flush_table()
                    bq_lines = []
                    while i < len(lines) and lines[i].strip().startswith(">"):
                        bq_lines.append(lines[i].strip()[1:].strip())
                        i += 1
                    _add_blockquote(doc, "\n".join(bq_lines))
                    continue

                _sep = line.strip()
                if _sep.startswith("|") and re.match(r"^[\|\-\:\s]+$", _sep):
                    i += 1
                    continue

                trow = _parse_table_row(line)
                if trow is not None:
                    pending_table_rows.append(trow)
                    i += 1
                    continue
                else:
                    _flush_table()

                stripped = line.strip()

                hm = re.match(r"^(#{1,3})\s+(.+)", stripped)
                if hm:
                    level = len(hm.group(1))
                    heading_text = hm.group(2).strip()
                    h = doc.add_heading(heading_text, level=min(level, 2))
                    for run_ in h.runs:
                        if level == 1:
                            run_.font.color.rgb = C_PURPLE
                        else:
                            run_.font.color.rgb = C_BLUE
                        run_.font.name = "微软雅黑"
                        run_.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
                    i += 1
                    continue

                bm = re.match(r"^([●\-\*•]\s*)(.+)", stripped)
                if bm:
                    bullet_char = bm.group(1).strip()
                    item_text = bm.group(2).strip()
                    p = doc.add_paragraph(style="List Bullet")
                    if p.runs:
                        p.runs[0].text = ""
                    line_color = None
                    if "✅" in item_text:
                        line_color = C_GREEN
                    elif "⚠️" in item_text or "⚠" in item_text:
                        line_color = C_ORANGE
                    elif "✗" in item_text or "❌" in item_text:
                        line_color = C_RED
                    sym = bullet_char if bullet_char in ("●", "•", "◆", "▸") else "●"
                    _add_styled_run(p, sym + " ", bold=False, color=line_color or C_BLUE, size=Pt(11))
                    _render_inline_formatted(p, item_text, C_GREEN, C_ORANGE, C_RED, C_BLUE, C_GRAY)
                    i += 1
                    continue

                om = re.match(r"^(\d+[\.\、]\s*)(.+)", stripped)
                if om:
                    num_prefix = om.group(1).strip()
                    item_text = om.group(2).strip()
                    p = doc.add_paragraph(style="List Number")
                    if p.runs:
                        p.runs[0].text = ""
                    line_color = None
                    if "✅" in item_text:
                        line_color = C_GREEN
                    elif "⚠️" in item_text or "⚠" in item_text:
                        line_color = C_ORANGE
                    elif "✗" in item_text or "❌" in item_text:
                        line_color = C_RED
                    _add_styled_run(p, num_prefix + " ", bold=True, color=C_GRAY, size=Pt(11))
                    _render_inline_formatted(p, item_text, C_GREEN, C_ORANGE, C_RED, C_BLUE, C_GRAY)
                    i += 1
                    continue

                p = doc.add_paragraph()
                _render_inline_formatted(p, stripped, C_GREEN, C_ORANGE, C_RED, C_BLUE, C_GRAY)
                i += 1

            _flush_table()

        doc.save(out)
        return f"Word 文档已生成：{out}\n（已保存为 .docx 文件，可在下方下载卡片中下载）"

    except Exception as e:
        return f"Word 文档生成失败：{type(e).__name__}: {e}"


def _render_inline_formatted(paragraph, text, C_GREEN, C_ORANGE, C_RED, C_BLUE, C_GRAY):
    """在段落中渲染内联格式（**粗体** / *斜体* / `代码` / 彩色标记）。"""
    from docx.shared import Pt
    from docx.oxml.ns import qn

    pattern = r"(\*\*.+?\*\*|\*.+?\*|`.+?`)"
    parts = re.split(pattern, text)

    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            inner = part[2:-2]
            _add_styled_run_inline(paragraph, inner, bold=True)
        elif part.startswith("*") and part.endswith("*"):
            inner = part[1:-1]
            _add_styled_run_inline(paragraph, inner, italic=True)
        elif part.startswith("`") and part.endswith("`"):
            inner = part[1:-1]
            r = paragraph.add_run(inner)
            r.font.name = "Consolas"
            r.font.size = Pt(10)
            r.font.color.rgb = C_GRAY
            r.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        else:
            _add_styled_run_inline(paragraph, part)


def _add_styled_run_inline(paragraph, text, bold=False, italic=False, override_color=None):
    """内联渲染辅助函数：处理带彩色标记的普通文本片段。"""
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.shared import RGBColor

    C_GREEN = RGBColor(0x16, 0xA3, 0x4A)
    C_ORANGE = RGBColor(0xF5, 0x9E, 0x0B)
    C_RED = RGBColor(0xEF, 0x44, 0x44)
    C_BLUE = RGBColor(0x3B, 0x82, 0xF6)
    C_PURPLE = RGBColor(0x8B, 0x5C, 0xF6)
    C_GRAY = RGBColor(0x6B, 0x72, 0x80)

    markers = [
        ("✅", C_GREEN), ("⚠️", C_ORANGE), ("⚠", C_ORANGE),
        ("✗", C_RED), ("❌", C_RED),
        ("🔵", C_BLUE), ("🟢", C_GREEN), ("🟡", C_ORANGE),
        ("🔴", C_RED), ("💡", C_ORANGE), ("📌", C_PURPLE),
    ]

    if not markers or not any(m[0] in text for m in markers):
        r = paragraph.add_run(text)
        r.bold = bold
        r.italic = italic
        r.font.name = "微软雅黑"
        r.font.size = Pt(11)
        r.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        if override_color:
            r.font.color.rgb = override_color
        return

    marker_chars = "".join(set(m[0] for m in markers))
    pat = "(" + "|".join(re.escape(m[0]) for m in markers) + ")"
    segments = re.split(pat, text)

    for seg in segments:
        if not seg:
            continue
        color = override_color
        for mc, mc_color in markers:
            if seg == mc:
                color = mc_color
                break
        r = paragraph.add_run(seg)
        r.bold = bold
        r.italic = italic
        r.font.name = "微软雅黑"
        r.font.size = Pt(11)
        r.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        if color:
            r.font.color.rgb = color
